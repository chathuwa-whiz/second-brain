"""
resume_extract.py — shared plain-text extraction for resume files, used by
both local_resumes.py (VPS/local filesystem) and drive_resumes.py (Google
Drive). Kept separate so neither storage backend needs to know about the
other's extraction details.
"""

import io

from docx import Document as DocxDocument
from pypdf import PdfReader


def extract_docx_text(raw: bytes) -> str:
    doc = DocxDocument(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def extract_pdf_text(raw: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text(raw: bytes, filename: str) -> str:
    """Dispatch on file extension. Raises ValueError for unsupported types."""
    lower = filename.lower()
    if lower.endswith(".docx"):
        return extract_docx_text(raw)
    if lower.endswith(".pdf"):
        return extract_pdf_text(raw)
    raise ValueError(f"Unsupported resume file type: {filename}")
