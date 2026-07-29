"""
drive_resumes.py — read resume files from a Google Drive folder via a
service account, and extract plain text from them.

Deliberately a SERVICE ACCOUNT integration, not OAuth user-consent: this
runs as an unattended background process (today on a local machine,
eventually on a headless VPS with no browser to complete a login flow), so
a service account with the folder explicitly shared to it is the right
fit — no login screen, no token refresh dance, just a JSON key file that
only has read access to whatever folder you explicitly share with it.

One-time setup (Google Cloud Console):
    1. Create/select a project, enable the "Google Drive API".
    2. APIs & Services -> Credentials -> Create Credentials -> Service
       Account. Any name is fine (e.g. "second-brain-resumes").
    3. On that service account, create a JSON key and download it.
       Treat this file like a password: never commit it, keep it out of
       the repo (it's covered by the same .gitignore pattern as .env).
    4. In Google Drive, share the resume folder with the service
       account's email (looks like
       second-brain-resumes@<project>.iam.gserviceaccount.com).
       Viewer access is enough — this tool never writes anything.
    5. Set GOOGLE_SERVICE_ACCOUNT_FILE (path to the JSON key) and
       RESUME_DRIVE_FOLDER_ID (the folder's ID, from its Drive URL) in
       job-tracker-mcp/.env.
"""

import io
import os
from pathlib import Path

from docx import Document as DocxDocument
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from pypdf import PdfReader

SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

# mimeType -> how we should handle it. Google Docs get exported to plain
# text via the API's export endpoint rather than downloaded as raw bytes.
SUPPORTED_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/pdf": "pdf",
    "application/vnd.google-apps.document": "gdoc",
}


def _get_drive_service():
    key_path = os.environ.get("GOOGLE_SERVICE_ACCOUNT_FILE")
    if not key_path or not Path(key_path).exists():
        raise RuntimeError(
            "GOOGLE_SERVICE_ACCOUNT_FILE is not set or the file doesn't "
            "exist. See the setup steps in this file's docstring."
        )
    creds = service_account.Credentials.from_service_account_file(key_path, scopes=SCOPES)
    return build("drive", "v3", credentials=creds)


def list_resume_files(folder_id: str) -> list[dict]:
    """List supported resume files (.docx / .pdf / Google Doc) directly
    inside a Drive folder. Returns [{"id", "name", "mimeType", ...}, ...].
    """
    service = _get_drive_service()
    mime_filter = " or ".join(f"mimeType='{m}'" for m in SUPPORTED_MIME_TYPES)
    query = f"'{folder_id}' in parents and trashed=false and ({mime_filter})"
    results = (
        service.files()
        .list(q=query, fields="files(id, name, mimeType, modifiedTime)")
        .execute()
    )
    return results.get("files", [])


def _extract_docx_text(raw: bytes) -> str:
    doc = DocxDocument(io.BytesIO(raw))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def _extract_pdf_text(raw: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def get_resume_text(file_id: str, mime_type: str) -> str:
    """Download a single Drive file and extract its plain text."""
    service = _get_drive_service()

    if mime_type == "application/vnd.google-apps.document":
        raw = service.files().export(fileId=file_id, mimeType="text/plain").execute()
        return raw.decode("utf-8") if isinstance(raw, bytes) else raw

    request = service.files().get_media(fileId=file_id)
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    raw = buffer.getvalue()

    if mime_type == "application/pdf":
        return _extract_pdf_text(raw)
    return _extract_docx_text(raw)
